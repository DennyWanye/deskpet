# SPDX-FileCopyrightText: 2026 DennyWanye
# SPDX-License-Identifier: BUSL-1.1

"""Word document tools — python-docx wrapper (beta builtin skill).

Three tools power the ``doc-edit`` skill:

* ``doc_create(spec)``        — build a brand-new ``.docx`` from a
                                 structured element list.
* ``doc_read(file_path)``     — read an existing ``.docx`` into a
                                 structured outline so the model can
                                 decide *where* to edit.
* ``doc_edit(file_path, ops)``— apply a list of targeted edit ops to an
                                 existing document, preserving styles.

Path policy: ``doc_read`` / ``doc_edit`` only accept a path the user has
picked through ``office_pick_file`` (see :mod:`office_paths`). ``doc_create``
writes to temp or a user-picked directory.

Edit philosophy — *targeted ops, never whole-doc rewrite*. We replace
text inside the existing run when possible (formatting fully preserved),
and only fall back to a paragraph-level rebuild (paragraph style kept,
run formatting collapsed to run[0]) when the match spans runs. This is
what keeps "把甲方名字改一下" from mangling the rest of a contract.

Failure philosophy: bad input → ``{"ok": false, "error": ...}``; a
single failing op is marked ``skipped`` and the rest still apply.
"""
from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any, Optional

from . import office_paths

log = logging.getLogger(__name__)

try:
    import docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _HAS_DOCX = True
except Exception:  # noqa: BLE001
    _HAS_DOCX = False


_ALIGN = {
    "left": "LEFT",
    "center": "CENTER",
    "right": "RIGHT",
    "justify": "JUSTIFY",
}


def _coerce(obj: Any) -> Optional[Any]:
    """Accept a Python object or a JSON string of one."""
    if isinstance(obj, str):
        try:
            return json.loads(obj)
        except json.JSONDecodeError:
            return None
    return obj


# ---------------------------------------------------------------------------
# Rich-formatting helpers (font color, mixed runs, lists, images, page #)
# ---------------------------------------------------------------------------
def _rgb(c: Any) -> Optional["RGBColor"]:
    """'#1F4E78' / '1F4E78' → RGBColor; bad input → None."""
    if not c:
        return None
    try:
        s = str(c).lstrip("#")
        if len(s) == 6:
            return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except (ValueError, TypeError):
        return None
    return None


def _apply_run_style(run, d: dict[str, Any]) -> None:
    """Apply bold/italic/underline/color/font_size/font_name to a run.
    Used both for whole-paragraph shorthand and per-run mixed formatting."""
    if d.get("bold"):
        run.bold = True
    if d.get("italic"):
        run.italic = True
    if d.get("underline"):
        run.underline = True
    size = d.get("font_size")
    if size:
        try:
            run.font.size = Pt(float(size))
        except (ValueError, TypeError):
            pass
    color = _rgb(d.get("color"))
    if color is not None:
        run.font.color.rgb = color
    name = d.get("font_name") or d.get("font")
    if name:
        run.font.name = str(name)


def _has_style(document, name: str) -> bool:
    try:
        _ = document.styles[name]
        return True
    except Exception:  # noqa: BLE001
        return False


def _add_list(document, el: dict[str, Any]) -> None:
    """Bullet / numbered list. ``el`` = {items:[str|{text,...}], ordered?,
    level?}. Falls back to plain paragraphs if the list style is missing."""
    items = el.get("items") or []
    if not isinstance(items, list):
        return
    ordered = bool(el.get("ordered"))
    base = "List Number" if ordered else "List Bullet"
    try:
        level = int(el.get("level", 0) or 0)
    except (ValueError, TypeError):
        level = 0
    name = base if level <= 0 else f"{base} {min(level + 1, 3)}"
    if not _has_style(document, name):
        name = base if _has_style(document, base) else ""
    for it in items:
        text = it.get("text", "") if isinstance(it, dict) else it
        para = (
            document.add_paragraph(style=name)
            if name
            else document.add_paragraph()
        )
        run = para.add_run(str(text))
        if isinstance(it, dict):
            _apply_run_style(run, it)


def _add_image(document, el: dict[str, Any]) -> None:
    """Insert a picture. ``el`` = {path|src, width_in?, align?}."""
    p = el.get("path") or el.get("src")
    if not p or not Path(str(p)).exists():
        log.warning("doc image path missing: %r", p)
        return
    kw: dict[str, Any] = {}
    w = el.get("width_in") or el.get("width_inches")
    if w:
        try:
            kw["width"] = Inches(float(w))
        except (ValueError, TypeError):
            pass
    try:
        document.add_picture(str(p), **kw)
    except Exception as exc:  # noqa: BLE001 — image is non-critical
        log.warning("doc image skipped: %s", exc)
        return
    align = _ALIGN.get(str(el.get("align", "")).lower())
    if align and document.paragraphs:
        document.paragraphs[-1].alignment = getattr(WD_ALIGN_PARAGRAPH, align)


def _shade_cell(cell, hex_fill: str) -> None:
    """Fill a table cell with a background color (python-docx has no API;
    inject ``<w:shd>`` into the cell's properties)."""
    s = str(hex_fill).lstrip("#").upper()
    if len(s) != 6:
        return
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), s)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_page_number_field(paragraph) -> None:
    """Inject a live { PAGE } field into ``paragraph`` (renders the real
    page number in Word/WPS; python-docx has no high-level API for it)."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _apply_header_footer(document, spec_obj: dict[str, Any]) -> None:
    """Top-level spec keys: header (str), footer (str), page_number (bool).
    Page number goes centered in the footer (its own line if a footer text
    is also set)."""
    header = spec_obj.get("header")
    footer = spec_obj.get("footer")
    page_number = bool(spec_obj.get("page_number"))
    if header is None and footer is None and not page_number:
        return
    for section in document.sections:
        if header is not None:
            section.header.is_linked_to_previous = False
            section.header.paragraphs[0].text = str(header)
        if footer is not None:
            section.footer.is_linked_to_previous = False
            section.footer.paragraphs[0].text = str(footer)
        if page_number:
            section.footer.is_linked_to_previous = False
            fp = section.footer.paragraphs[0]
            if footer is not None and fp.text:
                fp = section.footer.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_page_number_field(fp)


# ---------------------------------------------------------------------------
# doc_create
# ---------------------------------------------------------------------------
def _add_element(document, el: dict[str, Any]) -> None:
    # 兼容 LLM 常用的简写格式：element 无 "type" 字段，而是直接用
    # {heading:"...",level:N} / {paragraph:"..."} / {table:[...]} 作 key。
    # 归一到标准 {type,text} 后再走下面的渲染分支（否则 type 缺省 paragraph、
    # text 取不到 → 渲染出空段落，整篇 docx 正文为空）。
    if "type" not in el:
        for _k in ("heading", "paragraph", "table", "page_break", "list", "image"):
            if _k in el:
                el = dict(el)
                inner = el[_k]
                el["type"] = _k
                # LLM 实际常发**嵌套** dict 格式 {"heading":{"text":..,"level":..}}
                # / {"paragraph":{"text":..}} / {"table":{"rows":..,"header":..}}。
                # 此前把整个内层 dict 直接赋给 el["text"]，渲染出 str(dict)
                # 字面量（如 "{'text': '团队周报', 'level': 1}"）→ 正文全是 dict
                # 文本。修复：内层是 dict → 平铺其字段到 el；是标量 → 按旧简写
                # ({"heading":"文字"} / {"paragraph":"文字"} / {"table":[[...]]}) 取值。
                if isinstance(inner, dict):
                    for _ik, _iv in inner.items():
                        el.setdefault(_ik, _iv)
                elif _k in ("heading", "paragraph") and "text" not in el:
                    el["text"] = inner
                elif _k == "table" and "rows" not in el:
                    el["rows"] = inner
                elif _k == "list" and "items" not in el:
                    el["items"] = inner  # {"list":["a","b"]} 简写
                elif _k == "image" and "path" not in el:
                    el["path"] = inner  # {"image":"C:/a.png"} 简写
                break
    etype = (el.get("type") or "paragraph").lower()
    if etype == "heading":
        level = int(el.get("level", 1))
        level = max(0, min(level, 9))
        h = document.add_heading(str(el.get("text", "")), level=level)
        # 标题支持字色(品牌色标题),其余排版交给内置 Heading 样式。
        hc = _rgb(el.get("color"))
        if hc is not None:
            for run in h.runs:
                run.font.color.rgb = hc
    elif etype == "page_break":
        document.add_page_break()
    elif etype == "list":
        _add_list(document, el)
    elif etype == "image":
        _add_image(document, el)
    elif etype == "table":
        rows = el.get("rows") or []
        if not rows:
            return
        ncols = max((len(r) for r in rows), default=0)
        table = document.add_table(rows=len(rows), cols=ncols)
        table.style = str(el.get("style") or "Table Grid")
        # 表头底色(默认深蓝+白字),可用 header_fill / header_color 覆盖。
        hdr_fill = str(el.get("header_fill") or "1F4E78").lstrip("#")
        hdr_color = _rgb(el.get("header_color") or "FFFFFF")
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                val = row[ci] if ci < len(row) else ""
                cell = table.cell(ri, ci)
                cell.text = str(val)
                if ri == 0 and el.get("header"):
                    _shade_cell(cell, hdr_fill)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            if hdr_color is not None:
                                run.font.color.rgb = hdr_color
    else:  # paragraph
        para = document.add_paragraph()
        # 段内混排: runs=[{text,bold,color,...}] 让一段里多种格式并存
        # (如"这几个字红色 + 那几个字加粗")。无 runs 时退化为整段一种格式。
        runs = el.get("runs")
        if isinstance(runs, list) and runs:
            for rd in runs:
                if isinstance(rd, dict):
                    run = para.add_run(str(rd.get("text", "")))
                    _apply_run_style(run, rd)
                else:
                    para.add_run(str(rd))
        else:
            run = para.add_run(str(el.get("text", "")))
            _apply_run_style(run, el)  # bold/italic/underline/color/font_size
        align = _ALIGN.get(str(el.get("align", "")).lower())
        if align:
            para.alignment = getattr(WD_ALIGN_PARAGRAPH, align)


def doc_create(spec: Any, *, output_path: Optional[str] = None) -> dict[str, Any]:
    """Create a new .docx from ``spec`` = {title?, elements: [...]}.

    Element types: heading{text,level,color}, paragraph{text|runs,bold,
    italic,underline,color,align,font_size}, list{items,ordered,level},
    table{rows,header,header_fill,header_color,style}, image{path,width_in,
    align}, page_break. Top-level: header, footer, page_number.
    """
    if not _HAS_DOCX:
        return {"ok": False, "error": "python-docx not installed", "retriable": False}

    spec_obj = _coerce(spec)
    if not isinstance(spec_obj, dict):
        return {"ok": False, "error": "spec must be a dict or JSON object string", "retriable": False}

    elements = spec_obj.get("elements")
    if not isinstance(elements, list):
        return {"ok": False, "error": "spec.elements must be a list", "retriable": False}

    try:
        out_path = office_paths.resolve_for_write(
            output_path, default_prefix="deskpet-doc", default_suffix=".docx",
            default_kind="Doc",
        )
    except office_paths.PathError as exc:
        return {"ok": False, "error": str(exc), "retriable": False}

    try:
        document = Document()
        title = spec_obj.get("title")
        for el in elements:
            if isinstance(el, dict):
                _add_element(document, el)
        _apply_header_footer(document, spec_obj)
        if title:
            try:
                document.core_properties.title = str(title)
            except Exception:  # noqa: BLE001
                pass
        document.save(str(out_path))
    except Exception as exc:  # noqa: BLE001
        log.warning("doc_create failed: %s", exc, exc_info=True)
        return {"ok": False, "error": f"render failed: {exc}", "retriable": True}

    # WI-T1.2 D1：显式 emit artifacts[]（一等公民路径，保 BC）
    return {
        "ok": True,
        "path": str(out_path),
        "element_count": len(elements),
        "artifacts": [{
            "kind": "file",
            "path": str(out_path),
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "title": Path(str(out_path)).name,
        }],
    }


# ---------------------------------------------------------------------------
# doc_read
# ---------------------------------------------------------------------------
def doc_read(file_path: str) -> dict[str, Any]:
    """Read an existing .docx into a structured outline.

    Returns paragraphs (index, text, style) + tables (index, rows, cols).
    The path must have been authorized via the file picker.
    """
    if not _HAS_DOCX:
        return {"ok": False, "error": "python-docx not installed", "retriable": False}

    resolved = office_paths.resolve_for_read(file_path)
    if resolved is None:
        return {
            "ok": False,
            "error": (
                "file not authorized or not found — call office_pick_file "
                "first so the user can choose the document"
            ),
            "retriable": False,
        }

    try:
        document = Document(str(resolved))
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": f"cannot open document: {exc}", "retriable": False}

    paragraphs = []
    for i, para in enumerate(document.paragraphs):
        style = ""
        try:
            style = para.style.name if para.style else ""
        except Exception:  # noqa: BLE001
            pass
        paragraphs.append({"index": i, "text": para.text, "style": style})

    tables = []
    for i, table in enumerate(document.tables):
        tables.append({
            "index": i,
            "rows": len(table.rows),
            "cols": len(table.columns),
        })

    return {
        "ok": True,
        "path": str(resolved),
        "paragraphs": paragraphs,
        "tables": tables,
    }


# ---------------------------------------------------------------------------
# doc_edit
# ---------------------------------------------------------------------------
def _replace_in_paragraph(para, find: str, replace: str) -> bool:
    """Replace ``find`` with ``replace`` inside one paragraph.

    Run-level first (formatting fully preserved). If the match spans
    runs, rebuild: put the whole replaced text into run[0] (keeps run[0]
    formatting + the paragraph style) and clear the rest.
    """
    if not find or find not in para.text:
        return False
    # Fast path: match contained in a single run.
    for run in para.runs:
        if find in run.text:
            run.text = run.text.replace(find, replace)
            return True
    # Slow path: match spans runs — rebuild paragraph text.
    new_text = para.text.replace(find, replace)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)
    return True


def _apply_op(document, op: dict[str, Any]) -> dict[str, Any]:
    kind = (op.get("op") or "").lower()
    paragraphs = document.paragraphs

    if kind == "replace":
        find = str(op.get("find", ""))
        replace = str(op.get("replace", ""))
        idx = op.get("paragraph_index")
        changed = 0
        if idx is not None:
            try:
                target = paragraphs[int(idx)]
            except (IndexError, ValueError, TypeError):
                return {"op": kind, "status": "skipped", "reason": "bad paragraph_index"}
            if _replace_in_paragraph(target, find, replace):
                changed = 1
        else:
            for para in paragraphs:
                if _replace_in_paragraph(para, find, replace):
                    changed += 1
        if changed == 0:
            return {"op": kind, "status": "skipped", "reason": f"no match for {find!r}"}
        return {"op": kind, "status": "ok", "changed": changed}

    if kind == "set_paragraph_text":
        try:
            target = paragraphs[int(op.get("index"))]
        except (IndexError, ValueError, TypeError):
            return {"op": kind, "status": "skipped", "reason": "bad index"}
        text = str(op.get("text", ""))
        if target.runs:
            target.runs[0].text = text
            for run in target.runs[1:]:
                run.text = ""
        else:
            target.add_run(text)
        return {"op": kind, "status": "ok"}

    if kind == "insert_paragraph":
        text = str(op.get("text", ""))
        style = op.get("style")
        after = op.get("after_index")
        new_para = document.add_paragraph(text)
        if style:
            try:
                new_para.style = str(style)
            except Exception:  # noqa: BLE001
                pass
        if after is not None:
            try:
                anchor = paragraphs[int(after)]
                anchor._p.addnext(new_para._p)
            except (IndexError, ValueError, TypeError):
                return {"op": kind, "status": "ok", "note": "appended (bad after_index)"}
        return {"op": kind, "status": "ok"}

    if kind == "set_table_cell":
        try:
            table = document.tables[int(op.get("table_index"))]
            cell = table.cell(int(op.get("row")), int(op.get("col")))
        except (IndexError, ValueError, TypeError):
            return {"op": kind, "status": "skipped", "reason": "bad table coordinates"}
        cell.text = str(op.get("text", ""))
        return {"op": kind, "status": "ok"}

    return {"op": kind or "?", "status": "skipped", "reason": "unknown op"}


def doc_edit(file_path: str, ops: Any) -> dict[str, Any]:
    """Apply a list of edit ops to an existing .docx, in place.

    Ops: replace{find,replace,paragraph_index?}, set_paragraph_text{index,
    text}, insert_paragraph{text,after_index?,style?}, set_table_cell{
    table_index,row,col,text}.

    A failing op is recorded as ``skipped`` — the rest still apply.
    """
    if not _HAS_DOCX:
        return {"ok": False, "error": "python-docx not installed", "retriable": False}

    resolved = office_paths.resolve_for_read(file_path)
    if resolved is None:
        return {
            "ok": False,
            "error": (
                "file not authorized or not found — call office_pick_file "
                "first so the user can choose the document"
            ),
            "retriable": False,
        }

    ops_list = _coerce(ops)
    if ops_list is None:
        ops_list = []
    if not isinstance(ops_list, list):
        return {"ok": False, "error": "ops must be a list", "retriable": False}

    try:
        document = Document(str(resolved))
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "error": f"cannot open document: {exc}", "retriable": False}

    results = []
    for op in ops_list:
        if not isinstance(op, dict):
            results.append({"op": "?", "status": "skipped", "reason": "op not a dict"})
            continue
        try:
            results.append(_apply_op(document, op))
        except Exception as exc:  # noqa: BLE001 — one bad op never aborts the rest
            results.append({"op": op.get("op", "?"), "status": "skipped", "reason": str(exc)})

    applied = sum(1 for r in results if r.get("status") == "ok")
    if applied > 0:
        try:
            document.save(str(resolved))
        except Exception as exc:  # noqa: BLE001
            return {"ok": False, "error": f"save failed: {exc}", "retriable": True}

    return {
        "ok": True,
        "path": str(resolved),
        "applied": applied,
        "ops": results,
    }


# ---------------------------------------------------------------------------
# Tool registry wiring
# ---------------------------------------------------------------------------
_CREATE_SCHEMA = {
    "name": "doc_create",
    "description": (
        "Create a new Word .docx locally from a structured spec. Elements: "
        "heading{text,level,color}, paragraph{text OR runs:[{text,bold,italic,"
        "underline,color:'#RRGGBB',font_size}], bold,italic,color,align,"
        "font_size}, list{items,ordered,level}, table{rows,header,header_fill,"
        "header_color,style}, image{path,width_in,align}, page_break. "
        "Top-level spec keys: title, header (str), footer (str), page_number "
        "(bool, adds live page numbers). Saves to 桌宠/OutPut/Doc by default — "
        "tell the user the full returned path."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "spec": {"description": "Dict (or JSON string) {title?, elements: [...]}."},
            "output_path": {
                "type": "string",
                "description": "Absolute .docx path; omit to auto-create in temp.",
            },
        },
        "required": ["spec"],
    },
}

_READ_SCHEMA = {
    "name": "doc_read",
    "description": (
        "Read an existing .docx into a structured outline (paragraphs with "
        "index/text/style, tables with dimensions). Call this BEFORE doc_edit "
        "to learn where to edit. The file must first be chosen by the user "
        "via office_pick_file."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Absolute path to the .docx."},
        },
        "required": ["file_path"],
    },
}

_EDIT_SCHEMA = {
    "name": "doc_edit",
    "description": (
        "Apply targeted edits to an existing .docx in place, preserving "
        "styles. Ops: replace{find,replace,paragraph_index?}, "
        "set_paragraph_text{index,text}, insert_paragraph{text,after_index?,"
        "style?}, set_table_cell{table_index,row,col,text}. The file must "
        "first be chosen by the user via office_pick_file."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "file_path": {"type": "string", "description": "Absolute path to the .docx."},
            "ops": {"description": "List of edit-op dicts (or a JSON string of one)."},
        },
        "required": ["file_path", "ops"],
    },
}


def _handle_create(args: dict[str, Any], task_id: str = "") -> str:
    result = doc_create(
        args.get("spec"),
        output_path=(str(args["output_path"]) if args.get("output_path") else None),
    )
    return json.dumps(result, ensure_ascii=False)


def _handle_read(args: dict[str, Any], task_id: str = "") -> str:
    return json.dumps(doc_read(str(args.get("file_path", ""))), ensure_ascii=False)


def _handle_edit(args: dict[str, Any], task_id: str = "") -> str:
    result = doc_edit(str(args.get("file_path", "")), args.get("ops"))
    return json.dumps(result, ensure_ascii=False)


def _register() -> None:
    try:
        from .registry import registry

        registry.register("doc_create", "office", _CREATE_SCHEMA, _handle_create,
                           permission_category="write_file", timeout_seconds=30.0,
                           concurrency_safe=False)  # G3: writes .docx to disk
        registry.register("doc_read", "office", _READ_SCHEMA, _handle_read,
                           permission_category="read_file", timeout_seconds=20.0)
        registry.register("doc_edit", "office", _EDIT_SCHEMA, _handle_edit,
                           permission_category="write_file", timeout_seconds=30.0,
                           concurrency_safe=False)  # G3: mutates .docx on disk
    except Exception:  # noqa: BLE001
        pass


_register()

__all__ = ["doc_create", "doc_read", "doc_edit"]
