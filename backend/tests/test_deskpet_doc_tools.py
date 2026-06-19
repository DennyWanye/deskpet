# SPDX-FileCopyrightText: 2026 DennyWanye
# SPDX-License-Identifier: BUSL-1.1

"""TDD T2 — doc_tools.py (doc_create / doc_read / doc_edit)."""
from __future__ import annotations

from pathlib import Path

import pytest

from deskpet.tools import office_paths as op
from deskpet.tools import doc_tools as dt

pytestmark = pytest.mark.skipif(not dt._HAS_DOCX, reason="python-docx not installed")


@pytest.fixture(autouse=True)
def _auth(tmp_path: Path):
    op.clear_authorizations()
    op.authorize_path(tmp_path)
    yield
    op.clear_authorizations()


def _make_doc(tmp_path: Path, elements) -> str:
    out = tmp_path / "src.docx"
    r = dt.doc_create({"elements": elements}, output_path=str(out))
    assert r["ok"], r
    return r["path"]


def test_t2_1_create_basic(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "heading", "text": "标题", "level": 1},
        {"type": "paragraph", "text": "正文一段"},
    ])
    from docx import Document

    doc = Document(path)
    assert any("正文一段" in p.text for p in doc.paragraphs)


def test_t2_2_create_with_table(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "table", "rows": [["a", "b"], ["c", "d"]], "header": True},
    ])
    from docx import Document

    doc = Document(path)
    assert len(doc.tables) == 1
    assert len(doc.tables[0].rows) == 2


def test_t2_3_heading_levels(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "heading", "text": "H1", "level": 1},
        {"type": "heading", "text": "H2", "level": 2},
    ])
    from docx import Document

    doc = Document(path)
    styles = [p.style.name for p in doc.paragraphs]
    assert any("Heading 1" in s for s in styles)
    assert any("Heading 2" in s for s in styles)


def test_t2_4_doc_read_outline(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "heading", "text": "章节", "level": 1},
        {"type": "paragraph", "text": "段落"},
    ])
    r = dt.doc_read(path)
    assert r["ok"]
    assert len(r["paragraphs"]) >= 2
    assert all("index" in p and "text" in p and "style" in p for p in r["paragraphs"])


def test_t2_5_edit_replace(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "paragraph", "text": "甲方：旧名"},
        {"type": "paragraph", "text": "乙方：保持不变"},
    ])
    r = dt.doc_edit(path, [{"op": "replace", "find": "旧名", "replace": "新名"}])
    assert r["ok"] and r["applied"] == 1
    from docx import Document

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs]
    assert any("甲方：新名" in t for t in texts)
    assert any("乙方：保持不变" in t for t in texts)


def test_t2_6_edit_replace_scoped_to_index(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "paragraph", "text": "重复词"},
        {"type": "paragraph", "text": "重复词"},
    ])
    # paragraph_index targets the 2nd paragraph only.
    from docx import Document

    doc = Document(path)
    idx = next(i for i, p in enumerate(doc.paragraphs) if p.text == "重复词")
    second = next(i for i, p in enumerate(doc.paragraphs) if p.text == "重复词" and i > idx)
    r = dt.doc_edit(path, [{"op": "replace", "find": "重复词", "replace": "改了", "paragraph_index": second}])
    assert r["ok"] and r["applied"] == 1
    doc2 = Document(path)
    texts = [p.text for p in doc2.paragraphs]
    assert texts.count("重复词") == 1
    assert "改了" in texts


def test_t2_7_insert_paragraph(tmp_path: Path):
    path = _make_doc(tmp_path, [{"type": "paragraph", "text": "原段"}])
    from docx import Document

    before = len(Document(path).paragraphs)
    r = dt.doc_edit(path, [{"op": "insert_paragraph", "text": "新段", "after_index": 0}])
    assert r["ok"]
    after = len(Document(path).paragraphs)
    assert after == before + 1


def test_t2_8_set_table_cell(tmp_path: Path):
    path = _make_doc(tmp_path, [{"type": "table", "rows": [["a", "b"], ["c", "d"]]}])
    r = dt.doc_edit(path, [{"op": "set_table_cell", "table_index": 0, "row": 1, "col": 1, "text": "改了"}])
    assert r["ok"] and r["applied"] == 1
    from docx import Document

    assert Document(path).tables[0].cell(1, 1).text == "改了"


def test_t2_9_edit_preserves_style(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "heading", "text": "标题甲", "level": 1},
    ])
    dt.doc_edit(path, [{"op": "replace", "find": "甲", "replace": "乙"}])
    from docx import Document

    doc = Document(path)
    heading = next(p for p in doc.paragraphs if "标题" in p.text)
    assert "Heading 1" in heading.style.name


def test_t2_10_edit_unauthorized_path_rejected(tmp_path: Path):
    path = _make_doc(tmp_path, [{"type": "paragraph", "text": "x"}])
    op.clear_authorizations()  # revoke
    r = dt.doc_edit(path, [{"op": "replace", "find": "x", "replace": "y"}])
    assert not r["ok"] and r["retriable"] is False
    assert "office_pick_file" in r["error"]


def test_t2_11_read_missing_file():
    r = dt.doc_read("Z:\\nope\\missing.docx")
    assert not r["ok"]


def test_t2_12_edit_empty_ops_noop(tmp_path: Path):
    path = _make_doc(tmp_path, [{"type": "paragraph", "text": "x"}])
    r = dt.doc_edit(path, [])
    assert r["ok"] and r["applied"] == 0


def test_t2_13_failing_op_skipped_rest_apply(tmp_path: Path):
    path = _make_doc(tmp_path, [{"type": "paragraph", "text": "命中目标"}])
    r = dt.doc_edit(path, [
        {"op": "replace", "find": "不存在", "replace": "x"},
        {"op": "replace", "find": "命中", "replace": "已改"},
    ])
    assert r["ok"] and r["applied"] == 1
    assert r["ops"][0]["status"] == "skipped"
    assert r["ops"][1]["status"] == "ok"


# ─────────────────────────────────────────────────────────────────────
# 回归 (2026-06-04 工具层严测 TC-11)：LLM 实发**嵌套** element 格式
#   {"heading":{"text":..,"level":..}} / {"paragraph":{"text":..}}
# 旧 _add_element 归一化只认简写字符串({"heading":"文字"})，把整个内层
# dict 当 text → 正文渲染成字面量 "{'text': '标题', 'level': 1}"。
# 修复后内层 dict 应平铺其字段 → 正确提取 text/level。
# ─────────────────────────────────────────────────────────────────────
def test_t2_regression_nested_element_dict_format(tmp_path: Path):
    """嵌套 dict 格式 {heading:{text,level}} 必须提取 text，不能 str(dict)。"""
    path = _make_doc(tmp_path, [
        {"heading": {"text": "团队周报", "level": 1}},
        {"paragraph": {"text": "本周完成了核心开发。"}},
    ])
    from docx import Document

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 必须是干净文本，绝不能出现字面 dict（旧 bug 现象）
    assert "团队周报" in texts, texts
    assert any("本周完成了核心开发" in t for t in texts), texts
    assert not any("{'text'" in t or '{"text"' in t for t in texts), \
        f"BUG 回归：dict 被 str() 进正文 → {texts}"
    # heading 仍应用 Heading 样式、level 生效
    h = next(p for p in doc.paragraphs if p.text.strip() == "团队周报")
    assert "Heading" in h.style.name, h.style.name


def test_t2_regression_shorthand_string_still_works(tmp_path: Path):
    """旧简写格式 {heading:"文字"} / {paragraph:"文字"} 向后兼容不破。"""
    path = _make_doc(tmp_path, [
        {"heading": "简写标题"},
        {"paragraph": "简写正文"},
    ])
    from docx import Document

    doc = Document(path)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "简写标题" in texts and "简写正文" in texts, texts
    assert not any("{" in t for t in texts), texts


# --- 复杂 Word 升级 (list / mixed runs / color / image / header-footer-page#) ---

def test_t2_14_bullet_list(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "list", "items": ["第一点", "第二点", "第三点"]},
    ])
    from docx import Document

    doc = Document(path)
    bullets = [p for p in doc.paragraphs if p.text in ("第一点", "第二点", "第三点")]
    assert len(bullets) == 3, [p.text for p in doc.paragraphs]
    assert any("List" in (p.style.name or "") for p in bullets), \
        [p.style.name for p in bullets]


def test_t2_15_numbered_list(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "list", "ordered": True, "items": ["步骤A", "步骤B"]},
    ])
    from docx import Document

    doc = Document(path)
    items = [p for p in doc.paragraphs if p.text in ("步骤A", "步骤B")]
    assert len(items) == 2
    assert any("Number" in (p.style.name or "") for p in items), \
        [p.style.name for p in items]


def test_t2_16_mixed_runs_and_color(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "paragraph", "runs": [
            {"text": "红色加粗", "bold": True, "color": "#FF0000"},
            {"text": "普通"},
        ]},
    ])
    from docx import Document

    doc = Document(path)
    para = next(p for p in doc.paragraphs if "红色加粗" in p.text)
    assert len(para.runs) >= 2
    r0 = para.runs[0]
    assert r0.bold is True
    assert str(r0.font.color.rgb) == "FF0000"
    # 第二个 run 不应继承加粗
    assert para.runs[1].bold in (None, False)


def test_t2_17_paragraph_color_shorthand(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "paragraph", "text": "蓝字", "color": "#0000FF", "font_size": 14},
    ])
    from docx import Document

    doc = Document(path)
    para = next(p for p in doc.paragraphs if p.text == "蓝字")
    assert str(para.runs[0].font.color.rgb) == "0000FF"


def test_t2_18_table_header_shaded(tmp_path: Path):
    path = _make_doc(tmp_path, [
        {"type": "table", "header": True,
         "rows": [["列1", "列2"], ["a", "b"]]},
    ])
    from docx import Document

    doc = Document(path)
    cell = doc.tables[0].cell(0, 0)
    # <w:shd w:fill="1F4E78"> 注入到表头单元格
    tcpr_xml = cell._tc.get_or_add_tcPr().xml
    assert "shd" in tcpr_xml and "1F4E78" in tcpr_xml


def test_t2_19_image_element(tmp_path: Path):
    from PIL import Image as _PILImage

    img_path = tmp_path / "pic.png"
    _PILImage.new("RGB", (20, 20), (0, 128, 0)).save(img_path)
    path = _make_doc(tmp_path, [
        {"type": "heading", "text": "带图文档", "level": 1},
        {"type": "image", "path": str(img_path), "width_in": 2.0},
    ])
    from docx import Document

    doc = Document(path)
    # 图片落进 document → inline_shapes 至少 1
    assert len(doc.inline_shapes) == 1


def test_t2_20_header_footer_pagenumber(tmp_path: Path):
    out = tmp_path / "hf.docx"
    r = dt.doc_create({
        "header": "公司机密",
        "footer": "第 X 页",
        "page_number": True,
        "elements": [{"type": "paragraph", "text": "正文"}],
    }, output_path=str(out))
    assert r["ok"], r
    from docx import Document

    doc = Document(r["path"])
    sec = doc.sections[0]
    assert sec.header.paragraphs[0].text == "公司机密"
    # 页码字段 { PAGE } 注入到页脚 XML
    footer_xml = sec.footer._element.xml
    assert "PAGE" in footer_xml


def test_t2_21_default_path_under_output_doc(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("DESKPET_USER_DATA_DIR", str(tmp_path))
    r = dt.doc_create({"elements": [{"type": "paragraph", "text": "x"}]})
    assert r["ok"], r
    p = Path(r["path"])
    assert p.exists()
    assert p.parent.name == "Doc"
    assert p.parent.parent.name == "OutPut"
    p.unlink(missing_ok=True)
